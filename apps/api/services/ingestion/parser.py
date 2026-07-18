import csv
import io
import logging
import os
from typing import Optional

import fitz  # PyMuPDF
from docx import Document

from core.exceptions import ValidationError, ExternalServiceError

logger = logging.getLogger("services.ingestion.parser")


def parse_pdf(file_bytes: bytes) -> str:
    """
    Parses a PDF using PyMuPDF (fitz) and extracts all page text.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = []
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text.append(page_text)
        return "\n".join(text)
    except Exception as e:
        logger.error(f"PyMuPDF failed to parse PDF: {str(e)}")
        raise ExternalServiceError(f"Failed to parse PDF document: {str(e)}")


def parse_docx(file_bytes: bytes) -> str:
    """
    Parses a DOCX using python-docx and extracts paragraph text.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        # Also parse tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
        return "\n".join(text)
    except Exception as e:
        logger.error(f"python-docx failed to parse DOCX: {str(e)}")
        raise ExternalServiceError(f"Failed to parse DOCX document: {str(e)}")


def parse_csv(file_bytes: bytes) -> str:
    """
    Parses a CSV file and formats each row as an search-friendly, key-value representation.
    """
    try:
        # Decode file bytes
        content = file_bytes.decode("utf-8", errors="ignore")
        csv_file = io.StringIO(content)
        reader = csv.reader(csv_file)
        rows = list(reader)
        if not rows:
            return ""

        headers = [h.strip() for h in rows[0]]
        # Determine if first row is likely headers
        has_headers = len(rows) > 1 and all(h for h in headers)

        result_lines = []
        if has_headers:
            for idx, row in enumerate(rows[1:], start=1):
                row_items = []
                for col_idx, val in enumerate(row):
                    header_name = headers[col_idx] if col_idx < len(headers) else f"Column{col_idx+1}"
                    row_items.append(f"{header_name}: {val.strip()}")
                result_lines.append(f"Row {idx}: " + ", ".join(row_items))
        else:
            for idx, row in enumerate(rows, start=1):
                row_items = [val.strip() for val in row if val.strip()]
                if row_items:
                    result_lines.append(f"Row {idx}: " + ", ".join(row_items))

        return "\n".join(result_lines)
    except Exception as e:
        logger.error(f"CSV parser failed: {str(e)}")
        raise ExternalServiceError(f"Failed to parse CSV document: {str(e)}")


def parse_txt_or_md(file_bytes: bytes) -> str:
    """
    Decodes plain text/markdown file bytes.
    """
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"Text file decoding failed: {str(e)}")
        raise ExternalServiceError(f"Failed to decode plain text document: {str(e)}")


def parse_document(
    file_bytes: Optional[bytes] = None,
    file_path: Optional[str] = None,
    file_name: Optional[str] = None
) -> str:
    """
    Main entry point for parsing document files. Dispatches to the appropriate parser based on file type.
    Must provide either file_bytes (with file_name to detect extension) or file_path.
    """
    if file_bytes is None and file_path is None:
        raise ValidationError("Either file_bytes or file_path must be provided to parse_document")

    # Read from path if bytes are not provided
    if file_bytes is None and file_path is not None:
        try:
            if not os.path.exists(file_path):
                raise ValidationError(f"File not found at path: {file_path}")
            with open(file_path, "rb") as f:
                file_bytes = f.read()
            if not file_name:
                file_name = os.path.basename(file_path)
        except ValidationError:
            raise
        except Exception as e:
            logger.error(f"Failed to read file from path {file_path}: {str(e)}")
            raise ExternalServiceError(f"Failed to read file from disk: {str(e)}")

    assert file_bytes is not None

    # Determine extension
    name = (file_name or file_path or "").lower()
    _, ext = os.path.splitext(name)

    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes)
    elif ext == ".csv":
        return parse_csv(file_bytes)
    elif ext in (".txt", ".md", ".markdown", ".text"):
        return parse_txt_or_md(file_bytes)
    else:
        logger.warning(f"Unsupported file extension: {ext}, attempting plain text parsing")
        return parse_txt_or_md(file_bytes)
